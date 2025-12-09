"""
Service for document synthesis - creating unified documents from multiple sources.
"""
from typing import List, Dict, Any, Optional
from pathlib import Path
import json
import re
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.storage_service import StorageService
from app.services.openai_service import OpenAIService
from app.services.vector_store import VectorStore
from app.parsers import ParserFactory
from app.core.config import settings
from app.core.logging import get_logger

logger = get_logger(__name__)


class SynthesisService:
    """Service for synthesizing documents from multiple sources."""

    def __init__(self):
        """Initialize synthesis service."""
        self.storage_service = StorageService()
        self.openai_service = OpenAIService()
        self.vector_store = VectorStore(dimension=1536)
        self.parser_factory = ParserFactory()

    async def analyze_document_structures(
        self,
        filenames: List[str]
    ) -> Dict[str, Any]:
        """
        Analyze document structures from multiple PDF files and extract common structure.
        
        Args:
            filenames: List of PDF filenames to analyze
            
        Returns:
            Dictionary with structure analysis results
        """
        logger.info(f"Analyzing document structures for {len(filenames)} documents")
        
        data_dir = Path(settings.data_directory)
        structures = []
        all_sections = set()
        
        # Parse each document and extract structure
        for filename in filenames:
            file_path = data_dir / filename
            if not file_path.exists():
                logger.warning(f"File not found: {file_path}")
                continue
                
            try:
                # Read and parse document
                # Use LLM chunking to match the chunks in the vector store
                file_content = file_path.read_bytes()
                parser = self.parser_factory.get_parser(filename)
                # Pass openai_service to parser for LLM chunking (same as indexing)
                parser.openai_service = self.openai_service
                parsed_doc = await parser.parse(file_content, filename)
                
                # Extract structure (headings, sections)
                structure = self._extract_structure(parsed_doc)
                structure['filename'] = filename
                structure['page_count'] = parsed_doc.page_count
                structures.append(structure)
                
                # Collect all unique sections
                for section in structure.get('sections', []):
                    all_sections.add(section.get('title', ''))
                    
            except Exception as e:
                logger.error(f"Error analyzing {filename}: {e}")
                continue
        
        # Generate common structure using AI
        common_structure = await self._generate_common_structure(
            structures, list(all_sections)
        )
        
        return {
            'document_structures': structures,
            'common_structure': common_structure,
            'all_sections': sorted(list(all_sections))
        }
    
    def _extract_structure(self, parsed_doc) -> Dict[str, Any]:
        """
        Extract document structure from parsed document with improved heading detection.
        
        Args:
            parsed_doc: ParsedDocument object
            
        Returns:
            Dictionary with structure information
        """
        sections = []
        current_section = None
        
        # Analyze chunks to find sections - look for headings at start of chunks
        for chunk in parsed_doc.chunks:
            content = chunk.content.strip()
            
            # Split content into lines to find headings
            lines = content.split('\n')
            first_line = lines[0].strip() if lines else ""
            rest_content = '\n'.join(lines[1:]).strip() if len(lines) > 1 else ""
            
            # Check if first line is a heading
            if first_line and self._is_heading(first_line):
                # Save previous section
                if current_section:
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    'title': first_line[:200],  # Limit title length
                    'page_number': chunk.page_number,
                    'chunks': []
                }
                
                # Add rest of content if any
                if rest_content:
                    current_section['chunks'].append({
                        'content': rest_content,
                        'chunk_index': chunk.chunk_index,
                        'page_number': chunk.page_number
                    })
            elif current_section:
                # Add to current section
                current_section['chunks'].append({
                    'content': content,
                    'chunk_index': chunk.chunk_index,
                    'page_number': chunk.page_number
                })
            else:
                # No current section, check if this chunk might be a heading
                # (for documents that start with a heading)
                if self._is_heading(content) and len(content) < 200:
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        'title': content[:200],
                        'page_number': chunk.page_number,
                        'chunks': []
                    }
                else:
                    # Start a section with this content
                    current_section = {
                        'title': 'Content',  # Temporary title
                        'page_number': chunk.page_number,
                        'chunks': [{
                            'content': content,
                            'chunk_index': chunk.chunk_index,
                            'page_number': chunk.page_number
                        }]
                    }
        
        # Add last section
        if current_section:
            sections.append(current_section)
        
        # If no sections found or only generic sections, try to infer structure from content
        if not sections or all(s.get('title', '').lower() in ['content', 'main content'] for s in sections):
            # Try to find headings in the full text
            full_text = parsed_doc.full_text or '\n'.join([chunk.content for chunk in parsed_doc.chunks])
            inferred_sections = self._infer_sections_from_text(full_text, parsed_doc.chunks)
            if inferred_sections:
                sections = inferred_sections
        
        # Ensure we have at least one section
        if not sections:
            sections.append({
                'title': 'Main Content',
                'page_number': 1,
                'chunks': [
                    {
                        'content': chunk.content,
                        'chunk_index': chunk.chunk_index,
                        'page_number': chunk.page_number
                    }
                    for chunk in parsed_doc.chunks
                ]
            })
        
        return {
            'sections': sections,
            'total_chunks': len(parsed_doc.chunks),
            'total_pages': parsed_doc.page_count
        }
    
    def _infer_sections_from_text(self, text: str, chunks: Any) -> List[Dict]:
        """
        Try to infer document sections from text patterns.
        
        Args:
            text: Full document text
            chunks: List of document chunks
            
        Returns:
            List of inferred sections
        """
        sections = []
        lines = text.split('\n')
        
        current_section = None
        current_chunk_idx = 0
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line looks like a heading
            if self._is_heading(line) and len(line) < 200:
                # Save previous section
                if current_section and current_section.get('chunks'):
                    sections.append(current_section)
                
                # Find which chunk this might belong to
                page_num = 1
                if current_chunk_idx < len(chunks):
                    page_num = chunks[current_chunk_idx].page_number or 1
                
                current_section = {
                    'title': line[:200],
                    'page_number': page_num,
                    'chunks': []
                }
        
        # Add last section
        if current_section and current_section.get('chunks'):
            sections.append(current_section)
        
        return sections
    
    def _is_heading(self, text: str) -> bool:
        """
        Heuristic to detect if a text line is a heading.
        
        Args:
            text: Text to check
            
        Returns:
            True if likely a heading
        """
        # Remove extra whitespace
        text = text.strip()
        
        # Too long to be a heading
        if len(text) > 200:
            return False
        
        # Check for common heading patterns
        heading_patterns = [
            r'^\d+\.?\s+[A-Z]',  # Numbered headings: "1. Introduction"
            r'^[A-Z][A-Z\s]{2,50}$',  # All caps short lines
            r'^[A-Z][a-z]+\s+[A-Z]',  # Title Case
            r'^Chapter\s+\d+',  # Chapter headings
            r'^Section\s+\d+',  # Section headings
            r'^\d+\.\d+',  # Numbered sections: "1.1", "2.3"
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                return True
        
        # Check if it's a short line with no sentence-ending punctuation
        if len(text) < 100 and not text.endswith(('.', '!', '?')):
            # Check if it's mostly uppercase or title case
            if text.isupper() or (text.istitle() and len(text.split()) < 10):
                return True
        
        return False
    
    async def _generate_common_structure(
        self,
        structures: List[Dict],
        all_sections: List[str]
    ) -> Dict[str, Any]:
        """
        Generate a common structure that synthesizes all document structures using AI analysis.
        
        Args:
            structures: List of document structures
            all_sections: List of all unique section titles
            
        Returns:
            Common structure with inventory table
        """
        # Prepare detailed analysis for AI
        structure_summary = []
        for struct in structures:
            sections_info = []
            for section in struct.get('sections', []):
                # Get sample content from first chunk to understand the section
                chunks = section.get('chunks', [])
                sample_content = ""
                if chunks:
                    # Get first 500 chars of first chunk as sample
                    first_chunk = chunks[0].get('content', '')
                    sample_content = first_chunk[:500] if len(first_chunk) > 500 else first_chunk
                
                sections_info.append({
                    'title': section.get('title', ''),
                    'page_number': section.get('page_number'),
                    'chunk_count': len(chunks),
                    'sample_content': sample_content
                })
            
            structure_summary.append({
                'filename': struct.get('filename', ''),
                'total_pages': struct.get('total_pages', 0),
                'total_chunks': struct.get('total_chunks', 0),
                'sections': sections_info
            })
        
        # Create comprehensive prompt for AI
        prompt = f"""You are an expert technical document analyst. Analyze the following document structures from multiple source documents and create a comprehensive, unified table of contents (inventory table) that synthesizes all information.

DOCUMENT STRUCTURES TO ANALYZE:
{json.dumps(structure_summary, indent=2, ensure_ascii=False)}

YOUR TASK:
1. Analyze each document's structure and content samples to understand the topics and themes
2. Identify common themes, topics, and sections across all documents
3. Create a logical, hierarchical table of contents that:
   - Combines all unique topics from all documents
   - Groups related topics together
   - Creates a clear hierarchy (main sections, subsections, sub-subsections)
   - Uses descriptive, professional section titles
   - Organizes content in a logical flow (e.g., Introduction → Overview → Procedures → Safety → Maintenance → Appendices)
   - Includes sections that appear in only one document if they contain important information
   - Merges similar sections from different documents into unified sections

4. The structure should be comprehensive and cover all important topics from all documents

OUTPUT FORMAT:
Return a JSON object with a "sections" array. Each section object must have:
- "title": A clear, descriptive section title (string)
- "level": Hierarchy level - 1 for main sections, 2 for subsections, 3 for sub-subsections (integer)
- "order": Display order starting from 1 (integer)

EXAMPLE OUTPUT STRUCTURE:
{{
  "sections": [
    {{"title": "Introduction", "level": 1, "order": 1}},
    {{"title": "Overview", "level": 1, "order": 2}},
    {{"title": "General Information", "level": 2, "order": 3}},
    {{"title": "Scope", "level": 2, "order": 4}},
    {{"title": "Procedures", "level": 1, "order": 5}},
    {{"title": "Safety Requirements", "level": 1, "order": 6}},
    {{"title": "Maintenance", "level": 1, "order": 7}},
    {{"title": "Appendices", "level": 1, "order": 8}}
  ]
}}

IMPORTANT:
- Create a comprehensive structure that covers ALL topics from ALL documents
- Use logical, professional section titles
- Ensure proper hierarchy (main sections at level 1, subsections at level 2, etc.)
- Return ONLY valid JSON, no explanatory text before or after"""

        system_message = """You are an expert technical writer and document analyst with deep expertise in:
- Analyzing technical documents and extracting their structure
- Creating unified document structures from multiple sources
- Organizing information in logical, hierarchical formats
- Understanding technical content and identifying key topics and themes

Your task is to analyze multiple document structures and create a comprehensive, well-organized table of contents that synthesizes all information from all source documents."""

        try:
            logger.info("Generating inventory table using AI analysis...")
            response = await self.openai_service.generate_completion(
                prompt=prompt,
                system_message=system_message,
                temperature=0.4,  # Slightly higher for more creative structure
                max_tokens=4000,  # Allow for longer responses
                response_format={"type": "json_object"}
            )
            
            # Parse response
            result = json.loads(response)
            
            # Extract sections array
            if isinstance(result, dict):
                if 'sections' in result:
                    inventory_table = result['sections']
                elif 'table_of_contents' in result:
                    inventory_table = result['table_of_contents']
                elif 'inventory_table' in result:
                    inventory_table = result['inventory_table']
                else:
                    # Try to find any array in the response
                    for key, value in result.items():
                        if isinstance(value, list):
                            inventory_table = value
                            break
                    else:
                        inventory_table = []
            else:
                inventory_table = result if isinstance(result, list) else []
            
            # Ensure proper format and validate
            if not isinstance(inventory_table, list):
                logger.warning("AI response did not contain a list, using fallback")
                inventory_table = []
            
            # Clean and validate each section
            cleaned_table = []
            for idx, section in enumerate(inventory_table, 1):
                if isinstance(section, dict):
                    # Ensure required fields
                    title = section.get('title', '').strip()
                    if not title:
                        continue  # Skip empty titles
                    
                    level = int(section.get('level', 1))
                    if level < 1:
                        level = 1
                    if level > 4:
                        level = 4  # Limit to 4 levels
                    
                    order = int(section.get('order', idx))
                    
                    cleaned_table.append({
                        'title': title,
                        'level': level,
                        'order': order
                    })
                elif isinstance(section, str):
                    # Handle string entries
                    cleaned_table.append({
                        'title': section.strip(),
                        'level': 1,
                        'order': idx
                    })
            
            # Sort by order
            cleaned_table.sort(key=lambda x: x.get('order', 999))
            
            # Re-number orders to be sequential
            for idx, section in enumerate(cleaned_table, 1):
                section['order'] = idx
            
            logger.info(f"Generated inventory table with {len(cleaned_table)} sections")
            
            return {
                'inventory_table': cleaned_table,
                'total_sections': len(cleaned_table)
            }
            
        except Exception as e:
            logger.error(f"Error generating common structure: {e}", exc_info=True)
            # Fallback: create structure from all sections
            inventory_table = [
                {'title': section, 'level': 1, 'order': idx + 1}
                for idx, section in enumerate(sorted(all_sections)) if section.strip()
            ]
            return {
                'inventory_table': inventory_table,
                'total_sections': len(inventory_table)
            }
    
    async def find_paragraphs_for_section(
        self,
        section_title: str,
        filenames: List[str],
        top_k: int = 10,
        used_paragraph_ids: Optional[set] = None
    ) -> List[Dict[str, Any]]:
        """
        Find relevant paragraphs for a specific section using vector search and LLM validation.
        
        Args:
            section_title: Title of the section
            filenames: List of source document filenames
            top_k: Number of results to return
            used_paragraph_ids: Set of paragraph IDs already used in this session (to avoid duplicates)
            
        Returns:
            List of relevant paragraphs with metadata, validated by LLM
        """
        logger.info(f"Finding paragraphs for section: {section_title}")
        
        # Initialize used paragraph IDs set
        if used_paragraph_ids is None:
            used_paragraph_ids = set()
        
        # Use vector search to find relevant content
        query_embedding = await self.openai_service.generate_embedding(section_title)
        
        # Search with more results to have options after filtering
        search_results = await self.vector_store.search(
            query_vector=query_embedding,
            top_k=top_k * 3 * len(filenames),  # Get more results for filtering
            filters=None
        )
        
        # Filter by filename and remove duplicates
        candidate_paragraphs = []
        seen_chunk_ids = set()
        
        for result in search_results:
            filename = result.get('filename', '')
            if filename not in filenames:
                continue
            
            chunk_id = result.get('id', '')
            
            # Skip if already used in this session
            if chunk_id in used_paragraph_ids:
                continue
            
            # Skip duplicates
            if chunk_id in seen_chunk_ids:
                continue
            
            seen_chunk_ids.add(chunk_id)
            
            candidate_paragraphs.append({
                'id': chunk_id,
                'content': result.get('content', ''),
                'filename': filename,
                'page_number': result.get('page_number'),
                'section_title': result.get('section_title'),
                'chunk_index': result.get('chunk_index'),
                'score': result.get('score', 0.0),
                'distance': result.get('distance', 1.0)
            })
            
            # Get enough candidates for LLM validation
            if len(candidate_paragraphs) >= top_k * 2:
                break
        
        if not candidate_paragraphs:
            logger.warning(f"No candidate paragraphs found for section: {section_title}")
            return []
        
        # Validate paragraphs with LLM - check relevance to section
        validated_paragraphs = await self._validate_paragraphs_with_llm(
            section_title=section_title,
            paragraphs=candidate_paragraphs,
            filenames=filenames,
            max_results=top_k
        )
        
        return validated_paragraphs
    
    async def _validate_paragraphs_with_llm(
        self,
        section_title: str,
        paragraphs: List[Dict[str, Any]],
        filenames: List[str],
        max_results: int = 10
    ) -> List[Dict[str, Any]]:
        """
        Validate paragraph relevance to section using LLM analysis.
        
        Args:
            section_title: Title of the section
            paragraphs: List of candidate paragraphs
            filenames: Source document filenames
            max_results: Maximum number of results to return
            
        Returns:
            List of validated, relevant paragraphs
        """
        logger.info(f"Validating {len(paragraphs)} paragraphs for section: {section_title}")
        
        # Get document context for each filename
        document_contexts = {}
        for filename in filenames:
            # Get all chunks from this document for context
            doc_paragraphs = [
                p for p in paragraphs if p.get('filename') == filename
            ]
            if doc_paragraphs:
                # Get surrounding context from vector store
                context_text = await self._get_document_context(filename, doc_paragraphs[0])
                document_contexts[filename] = context_text
        
        # Prepare paragraphs for LLM validation
        paragraphs_info = []
        for para in paragraphs[:max_results * 2]:  # Limit for token efficiency
            paragraphs_info.append({
                'id': para.get('id'),
                'content': para.get('content', '')[:800],  # Limit content length
                'filename': para.get('filename'),
                'page_number': para.get('page_number'),
                'score': para.get('score', 0.0)
            })
        
        # Create prompt for LLM validation
        prompt = f"""You are an expert document analyst. Your task is to evaluate which paragraphs are truly relevant for a specific section in a synthesis document.

SECTION TITLE: "{section_title}"

CANDIDATE PARAGRAPHS:
{json.dumps(paragraphs_info, indent=2, ensure_ascii=False)}

DOCUMENT CONTEXTS:
{json.dumps({k: v[:1000] for k, v in document_contexts.items()}, indent=2, ensure_ascii=False)}

YOUR TASK:
1. Analyze each paragraph's content in the context of the section title
2. Determine if the paragraph is truly relevant to the section topic
3. Consider the document context to understand the paragraph's meaning
4. Rank paragraphs by relevance (most relevant first)
5. Exclude paragraphs that are not relevant to the section topic

OUTPUT FORMAT:
Return a JSON object with a "validated_paragraphs" array. Each object should have:
- "id": The paragraph ID (string)
- "is_relevant": true if relevant, false if not (boolean)
- "relevance_score": A score from 0.0 to 1.0 indicating relevance (float)
- "reason": Brief explanation of why it's relevant or not (string)

Only include paragraphs that are relevant (is_relevant: true). Order them by relevance_score (highest first).

Return ONLY valid JSON, no other text."""

        system_message = """You are an expert technical document analyst specializing in content relevance assessment.
Your task is to carefully evaluate whether paragraphs are truly relevant to a given section topic, considering the full document context."""

        try:
            response = await self.openai_service.generate_completion(
                prompt=prompt,
                system_message=system_message,
                temperature=0.2,  # Lower temperature for more consistent evaluation
                max_tokens=3000,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response)
            
            # Extract validated paragraphs
            validated_list = result.get('validated_paragraphs', [])
            if not isinstance(validated_list, list):
                validated_list = []
            
            # Filter to only relevant paragraphs and sort by relevance score
            relevant_paragraphs = [
                vp for vp in validated_list 
                if isinstance(vp, dict) and vp.get('is_relevant', False)
            ]
            relevant_paragraphs.sort(key=lambda x: x.get('relevance_score', 0.0), reverse=True)
            
            # Map back to original paragraph data
            para_map = {p.get('id'): p for p in paragraphs}
            final_paragraphs = []
            
            for vp in relevant_paragraphs[:max_results]:
                para_id = vp.get('id')
                if para_id in para_map:
                    original_para = para_map[para_id].copy()
                    original_para['llm_relevance_score'] = vp.get('relevance_score', 0.0)
                    # Don't add llm_reason to keep response clean
                    final_paragraphs.append(original_para)
            
            logger.info(f"LLM validated {len(final_paragraphs)} relevant paragraphs out of {len(paragraphs)} candidates")
            
            return final_paragraphs
            
        except Exception as e:
            logger.error(f"Error validating paragraphs with LLM: {e}", exc_info=True)
            # Fallback: return top paragraphs by vector score
            return sorted(paragraphs, key=lambda x: x.get('score', 0.0), reverse=True)[:max_results]
    
    async def _get_document_context(self, filename: str, sample_paragraph: Dict[str, Any]) -> str:
        """
        Get document context around a paragraph for better LLM understanding.
        
        Args:
            filename: Document filename
            sample_paragraph: A sample paragraph from the document
            
        Returns:
            Context text from the document
        """
        try:
            # Get all paragraphs from this document
            doc_paragraphs = []
            for para_id, metadata in self.vector_store.metadata.items():
                if metadata.get('filename') == filename:
                    doc_paragraphs.append({
                        'content': metadata.get('content', ''),
                        'page_number': metadata.get('page_number'),
                        'chunk_index': metadata.get('chunk_index', 0)
                    })
            
            # Sort by chunk_index
            doc_paragraphs.sort(key=lambda x: x.get('chunk_index', 0))
            
            # Get context around the sample paragraph (surrounding chunks)
            sample_idx = sample_paragraph.get('chunk_index', 0)
            context_start = max(0, sample_idx - 3)
            context_end = min(len(doc_paragraphs), sample_idx + 4)
            
            context_paragraphs = doc_paragraphs[context_start:context_end]
            context_text = '\n\n'.join([p.get('content', '') for p in context_paragraphs])
            
            return context_text[:2000]  # Limit context length
            
        except Exception as e:
            logger.warning(f"Error getting document context: {e}")
            return ""
    
    async def generate_synthesis_document(
        self,
        inventory_table: List[Dict],
        selected_paragraphs: Dict[str, List[str]],  # section_title -> list of paragraph IDs
        filenames: List[str]
    ) -> bytes:
        """
        Generate the final synthesis document from selected paragraphs as DOCX.
        
        Args:
            inventory_table: Final inventory table structure
            selected_paragraphs: Dictionary mapping section titles to selected paragraph IDs
            filenames: Source document filenames
            
        Returns:
            Generated document as DOCX bytes
        """
        logger.info("Generating synthesis document as DOCX")
        
        # Load paragraph content from vector store
        vector_store = VectorStore(dimension=1536)
        # Access metadata directly from vector store
        metadata = vector_store.metadata
        
        # Create a new Document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('Synthesis Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add introduction paragraph
        intro = doc.add_paragraph('This document synthesizes information from the following sources:')
        intro_format = intro.paragraph_format
        intro_format.space_after = Pt(12)
        
        # Add source list
        for filename in filenames:
            source_para = doc.add_paragraph(filename, style='List Bullet')
            source_para.paragraph_format.space_after = Pt(6)
        
        # Add separator
        doc.add_paragraph('─' * 50).paragraph_format.space_after = Pt(12)
        
        # Build document section by section
        for section in sorted(inventory_table, key=lambda x: x.get('order', 999)):
            section_title = section.get('title', '')
            level = section.get('level', 1)
            
            # Add section header with appropriate heading level
            # DOCX supports heading levels 1-9, map our levels accordingly
            heading_level = min(level, 9)
            heading = doc.add_heading(section_title, level=heading_level)
            heading.paragraph_format.space_before = Pt(12) if level == 1 else Pt(6)
            heading.paragraph_format.space_after = Pt(6)
            
            # Add selected paragraphs for this section
            paragraph_ids = selected_paragraphs.get(section_title, [])
            
            if paragraph_ids:
                for para_id in paragraph_ids:
                    # Find paragraph in metadata
                    para_data = metadata.get(para_id, {})
                    content = para_data.get('content', '')
                    
                    if content:
                        # Add paragraph content
                        para = doc.add_paragraph(content)
                        para.paragraph_format.space_after = Pt(12)
                        para.paragraph_format.first_line_indent = Inches(0.25) if level > 1 else Inches(0)
                        
                        # Add source reference as italic
                        filename = para_data.get('filename', 'Unknown')
                        page = para_data.get('page_number', '?')
                        source_ref = doc.add_paragraph(f'[Source: {filename}, page {page}]', style='Intense Quote')
                        source_ref.paragraph_format.space_after = Pt(12)
                        # Make it italic
                        for run in source_ref.runs:
                            run.italic = True
                            run.font.size = Pt(9)
            else:
                # No content selected
                empty_para = doc.add_paragraph('[No content selected for this section]', style='Intense Quote')
                empty_para.paragraph_format.space_after = Pt(12)
                for run in empty_para.runs:
                    run.italic = True
                    run.font.size = Pt(10)
        
        # Save document to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info("Synthesis document generated successfully as DOCX")
        return doc_bytes.getvalue()

